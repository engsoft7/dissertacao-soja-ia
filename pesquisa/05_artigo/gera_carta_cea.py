# -*- coding: utf-8 -*-
"""Cover letter para a submissão na Computers and Electronics in Agriculture."""
import os

RAIZ = os.path.dirname(os.path.abspath(__file__))
SAIDA = os.path.join(RAIZ, 'saida')
os.makedirs(SAIDA, exist_ok=True)

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH as AL
from docx.shared import Cm, Pt

F = 'Times New Roman'
d = Document()
s = d.sections[0]
s.page_width, s.page_height = Cm(21.0), Cm(29.7)
s.top_margin = s.bottom_margin = Cm(2.5)
s.left_margin = s.right_margin = Cm(2.5)
d.styles['Normal'].font.name = F
d.styles['Normal'].font.size = Pt(11)


def p(t='', neg=False, align=AL.JUSTIFY, dep=10, lh=1.15):
    par = d.add_paragraph()
    par.alignment = align
    par.paragraph_format.line_spacing = lh
    par.paragraph_format.space_after = Pt(dep)
    if t:
        r = par.add_run(t)
        r.font.name, r.font.size, r.font.bold = F, Pt(11), neg
    return par


p('Maycon Lima dos Santos', align=AL.LEFT, dep=0)
p('Graduate Program in Applied Computing (PPCA)', align=AL.LEFT, dep=0)
p('Universidade Federal do Pará — Tucuruí, PA, Brazil', align=AL.LEFT, dep=0)
p('mayconlimasan@gmail.com', align=AL.LEFT, dep=16)

p('5 September 2026', align=AL.LEFT, dep=16)

p('To the Editors-in-Chief', align=AL.LEFT, dep=0)
p('Computers and Electronics in Agriculture', align=AL.LEFT, dep=16)

p('Dear Editors,')

p('I submit for your consideration the manuscript "Machine learning crop yield '
  'prediction is bounded by target-variable quality: a diagnostic of value repetition '
  'in official statistics and its effect on model evaluation", as a Research Paper.')

p('Machine learning for crop yield prediction is evaluated against official statistics '
  'taken as ground truth. Your journal has published the systematic reviews that map '
  'this field (van Klompenburg et al., 2020; Leukel et al., 2023) and the recent work '
  'that questions how much reported accuracy survives an explicit statistical reference '
  '(Wang et al., 2024). None of them offers a way to check whether the target variable '
  'carries the interannual variation the model is asked to predict. That is the gap this '
  'manuscript addresses.')

p('I propose a diagnostic — the rate at which consecutive values repeat — and apply it '
  'to the Brazilian Municipal Agricultural Production survey for soybean across 2,806 '
  'municipalities and 24 crop years. Three results follow. First, 17.6% of the 46,536 '
  'consecutive-season pairs report strictly identical yields, more than a hundred '
  'standard deviations above a conservative permutation null. Second, what predicts the '
  'rate is not the region but the scale of the crop in the municipality: the correlation '
  'with planted area is −0.456, and strengthens to −0.489 when the Amazonian states are '
  'excluded, so the phenomenon is not confined to agricultural frontiers. Third, where '
  'the rate is highest, four standard algorithms fed with climatic and spectral '
  'predictors fail to outperform a baseline built only from municipal history and '
  'temporal trend.')

p('On the fit to your scope, I would like to be explicit. The manuscript does not '
  'propose a new algorithm, and it uses established implementations. Its contribution is '
  'to the machine learning pipeline rather than to agronomy: the diagnostic is computed '
  'from the target series alone, before any model is fitted, costs a single pass over the '
  'data, and identifies datasets on which model comparison cannot be informative. It is '
  'also actionable, because the property it detects is predicted by an observable '
  'covariate that allows affected records to be identified in advance of modelling. I '
  'believe this addresses a methodological blind spot in the literature your journal '
  'publishes, and I would welcome your judgement on whether it is better placed here or '
  'in Smart Agricultural Technology.')

p('The dataset assembled for this study — 49,342 municipality-season records from the '
  'primary source, together with the derived measures and every analysis script — is '
  'openly archived at Zenodo under DOI 10.5281/zenodo.22343315 and available at '
  'https://github.com/engsoft7/dissertacao-soja-ia. Every numerical claim in the '
  'manuscript can be recomputed from that deposit.')

p('The work is original, has not been published elsewhere, and is not under '
  'consideration by another journal. It is single-authored. Part of the Pará case study '
  'derives from my master\'s dissertation at the Universidade Federal do Pará; the '
  'national analysis that forms the core of this manuscript is new and was not part of '
  'that work. I declare no competing interests. The use of generative AI in manuscript '
  'preparation is declared in the manuscript, as required by Elsevier policy.')

p('I thank you for your time and look forward to your response.')

p('Sincerely,', dep=16)
p('Maycon Lima dos Santos', align=AL.LEFT, dep=0)
p('Graduate Program in Applied Computing, Universidade Federal do Pará', align=AL.LEFT, dep=0)

d.save(os.path.join(SAIDA, 'Cover_letter.docx'))
print('gerada')

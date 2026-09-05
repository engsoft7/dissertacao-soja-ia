# -*- coding: utf-8 -*-
"""Manuscrito em inglês para a Computers and Electronics in Agriculture."""
import os

RAIZ = os.path.dirname(os.path.abspath(__file__))
SAIDA = os.path.join(RAIZ, 'saida')
os.makedirs(SAIDA, exist_ok=True)

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH as AL
from docx.shared import Cm, Pt
from docx.enum.table import WD_TABLE_ALIGNMENT

F = 'Times New Roman'


def novo_doc():
    d = Document()
    s = d.sections[0]
    s.page_width, s.page_height = Cm(21.0), Cm(29.7)
    s.top_margin = s.bottom_margin = Cm(2.5)
    s.left_margin = s.right_margin = Cm(2.5)
    d.styles['Normal'].font.name = F
    d.styles['Normal'].font.size = Pt(12)
    return d


d = novo_doc()


def par(t, tam=12, neg=False, ital=False, align=AL.JUSTIFY, lh=2.0, dep=0, doc=None):
    doc = doc or d
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.line_spacing = lh
    p.paragraph_format.space_after = Pt(dep)
    if t:
        r = p.add_run(t)
        r.font.name, r.font.size, r.font.bold, r.font.italic = F, Pt(tam), neg, ital
    return p


def sec(t):
    par(t, 12, True, align=AL.LEFT, dep=6)


def corpo(t):
    par(t)


def cap(t):
    par(t, 10, align=AL.LEFT, lh=1.0, dep=9)


def _borda_horizontal(t):
    """Só linhas horizontais: o guia proíbe réguas verticais e sombreamento."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tbl = t._tbl
    pr = tbl.tblPr
    for velha in pr.findall(qn('w:tblBorders')):
        pr.remove(velha)
    b = OxmlElement('w:tblBorders')
    for lado, val in [('top', 'single'), ('bottom', 'single'),
                      ('left', 'none'), ('right', 'none'),
                      ('insideH', 'single'), ('insideV', 'none')]:
        e = OxmlElement(f'w:{lado}')
        e.set(qn('w:val'), val)
        e.set(qn('w:sz'), '6')
        e.set(qn('w:color'), '000000')
        b.append(e)
    pr.append(b)


def tabela(cab, linhas, negrito_em=()):
    t = d.add_table(rows=1 + len(linhas), cols=len(cab))
    t.style = 'Table Grid'
    _borda_horizontal(t)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, c in enumerate(cab):
        cel = t.cell(0, j).paragraphs[0]
        cel.alignment = AL.CENTER
        cel.paragraph_format.line_spacing = 1.0
        r = cel.add_run(c)
        r.font.name, r.font.size, r.font.bold = F, Pt(10), True
    for i, ln in enumerate(linhas, 1):
        for j, c in enumerate(ln):
            cel = t.cell(i, j).paragraphs[0]
            cel.alignment = AL.LEFT if j == 0 else AL.CENTER
            cel.paragraph_format.line_spacing = 1.0
            r = cel.add_run(c)
            r.font.name, r.font.size = F, Pt(10)
            r.font.bold = (i - 1) in negrito_em
    return t


# ══════════════════════════ TITLE PAGE ══════════════════════════
par('Machine learning crop yield prediction is bounded by target-variable '
    'quality: a diagnostic of value repetition in official statistics and its '
    'effect on model evaluation', 14, True, align=AL.LEFT, lh=1.5, dep=12)

par('Maycon Lima dos Santos', 12, align=AL.LEFT, lh=1.0, dep=6)
par('Graduate Program in Applied Computing (PPCA), Universidade Federal do Pará, '
    'Rua Itaipu s/n, Vila Permanente, 68464-000, Tucuruí, PA, Brazil',
    11, align=AL.LEFT, lh=1.0, dep=6)
par('Corresponding author: Maycon Lima dos Santos — mayconlimasan@gmail.com',
    11, align=AL.LEFT, lh=1.0, dep=12)

# ══════════════════════════ ABSTRACT ══════════════════════════
sec('Abstract')
par('Machine learning pipelines for crop yield prediction are evaluated against '
    'official statistics taken as ground truth, yet the internal structure of that '
    'target variable is rarely examined and no standard diagnostic exists for it. '
    'We propose one — the rate at which consecutive values repeat — and apply it to '
    'the Brazilian Municipal Agricultural Production survey (PAM/IBGE) for soybean '
    'across 2,806 municipalities in every state from 2001 to 2024. Of 46,536 '
    'consecutive-season pairs, 17.6% report strictly identical values, against '
    '7.5% ± 0.1 expected under a conservative permutation null that preserves each '
    'series\' marginal distribution — more than a hundred standard deviations above '
    'chance, with none of 2,000 permutations reaching the observed rate. The rate '
    'ranges from 3.7% to 45.7% across states, and what predicts it is not the region '
    'but the scale of the crop: the Spearman correlation between municipal planted '
    'area and repetition rate is −0.459, strengthening to −0.503 when the Legal '
    'Amazon states are excluded. Where the rate is highest, four standard algorithms '
    'fed with climatic and spectral predictors failed to outperform a baseline using '
    'only municipal history and temporal trend (R² = 0.216 for both), which shows '
    'that reported model accuracy in this task is bounded by a property of the '
    'target variable rather than by model capacity. We recommend that the repetition '
    'rate be reported alongside the proportion of missing data whenever official '
    'statistics are used as ground truth.',
    12, align=AL.JUSTIFY, lh=2.0, dep=12)

sec('Keywords')
par('crop yield prediction; data quality; official statistics; machine learning; '
    'Brazil', 12, align=AL.LEFT, lh=1.0, dep=12)

# ══════════════════════════ 1 ══════════════════════════
sec('1. Introduction')
corpo('Early prediction of crop yield supports decisions in logistics, credit and '
      'public policy. Over the past decade, machine learning fed by climate series '
      'and remotely sensed vegetation indices has become the predominant instrument '
      'for this task, with reported performance of the order of 10–20% relative '
      'error at the municipal or county scale (van Klompenburg et al., 2020; Khaki '
      'et al., 2020).')
corpo('This literature shares an assumption that is rarely examined: that the '
      'official statistic used as the target variable is a reliable measure of the '
      'yield actually obtained. In Brazil this reference is the Municipal '
      'Agricultural Production survey (Produção Agrícola Municipal, PAM), compiled '
      'by the Brazilian Institute of Geography and Statistics (IBGE). The survey '
      'documentation states, however, that municipal values do not result from '
      'census measurement but from estimates produced by the field agent, drawing '
      'on contacts with producers and technicians and on local knowledge (IBGE, '
      '2024).')
corpo('That design is operationally adequate where the network of informants is '
      'dense, but its effects where the network is sparse remain poorly documented. '
      'The question is usually framed in regional terms — the problem is assumed to '
      'belong to agricultural frontiers — and it is that framing this study puts to '
      'the test, by measuring the structure of the series across the entire national '
      'territory rather than in a region chosen in advance.')
corpo('The contribution is methodological and addressed to the machine learning '
      'pipeline rather than to agronomy. Systematic reviews of this field published '
      'in this journal catalogue predictors, architectures and validation protocols '
      '(van Klompenburg et al., 2020; Leukel et al., 2023), and recent work has begun '
      'to question how much of the reported accuracy survives an explicit statistical '
      'reference (Wang et al., 2024). None of them provides a way to check whether the '
      'target variable carries the interannual variation the model is asked to '
      'predict. The repetition rate proposed here is such a check: it is computed from '
      'the target series alone, before any model is fitted, costs a single pass over '
      'the data, and identifies datasets on which model comparison cannot be '
      'informative. We also show that the diagnostic is actionable, since the property '
      'it detects is predicted by an observable covariate — the scale of the crop in '
      'the municipality — which allows the affected records to be identified in '
      'advance of modelling.')
corpo('Four questions guide the investigation: (i) how frequently yield values '
      'repeat between consecutive seasons in the PAM; (ii) whether that frequency is '
      'compatible with expected agronomic variability; (iii) what predicts it — the '
      'region or the scale of the crop; and (iv) what consequence it imposes on '
      'machine learning models calibrated on that base. The contribution is to '
      'quantify a structural limit not previously reported, to identify its '
      'predictor, and to offer a diagnostic procedure that is replicable for other '
      'crops and territories.')

# ══════════════════════════ 2 ══════════════════════════
sec('2. Related work')
corpo('Machine learning for crop yield prediction consolidated as a field of its own '
      'over the past decade. The systematic review of van Klompenburg et al. (2020) '
      'identifies neural networks and tree-based methods as the predominant families '
      'and the coefficient of determination as the most reported metric. Leukel et '
      'al. (2023), reviewing early prediction at field scale, observe accuracy '
      'increasing over time but accompanied by marked heterogeneity in validation '
      'protocols. In both, methodological discussion concentrates on predictors, '
      'architectures and validation protocols; the quality of the target variable '
      'itself is not among the axes of analysis.')
corpo('The dominant methodological lineage originates in work applying deep learning '
      'to remote sensing series at county scale. You et al. (2017) introduced the '
      'combination of convolutional networks with Gaussian processes over spectral '
      'histograms; Khaki and Wang (2019) and Khaki et al. (2020) extended the '
      'approach with recurrent architectures; and Sun et al. (2019) applied a '
      'CNN-LSTM model to soybean in United States counties. Maimaitijiang et al. '
      '(2020) demonstrated gains from multimodal fusion at plot scale. Recent work '
      'densifies this line, combining multiple orbital sources in residual, '
      'recurrent and graph architectures (Fathi et al., 2025; Fu et al., 2025; '
      'Ingole et al., 2025).')
corpo('In Brazil, the literature concentrates on consolidated producing regions. '
      'Richetti et al. (2018) estimated soybean yield in Paraná from '
      'phenology-adjusted vegetation indices; Schwalbert et al. (2020) integrated '
      'meteorological and orbital data in the south of the country; Barbosa dos '
      'Santos et al. (2021) employed artificial neural networks; and von Bloh et al. '
      '(2023) built a municipal reference base for seven producing states. In all of '
      'them the PAM is taken as ground truth.')
corpo('Two studies approach the problem addressed here from a different angle. Li et '
      'al. (2024) show that removing the global trend from yield series markedly '
      'improves the accuracy of XGBoost-based models, evidence that a substantial '
      'share of the variance explained in such studies comes from the trend '
      'component rather than from the interannual response to environmental '
      'conditions. Wang et al. (2024) compare deep learning with panel regression '
      'for Argentine soybean, being among the few to explicitly confront complex '
      'models with a simple statistical reference. Both signal that the gain '
      'attributed to environmental variables may be smaller than the aggregate '
      'metric suggests.')
corpo('The gap this study addresses lies one step earlier. The cited studies evaluate '
      'their models against the official statistic without examining whether it '
      'carries, itself, the interannual variation they set out to predict. '
      'Investigating the structure of the series before modelling it is a '
      'requirement of validity, not an accessory step.')

# ══════════════════════════ 3 ══════════════════════════
sec('3. Materials and methods')

sec('3.1. The national dataset')
corpo('The repetition measure was conducted on the primary source: table 5457 of the '
      'SIDRA/IBGE system, from which municipal average yield and planted area of '
      'soybean were extracted for every Brazilian municipality, for the 2001–2024 '
      'crop years. The resulting dataset comprises 49,342 municipality-season records '
      'from 2,806 municipalities across all states. Of these, 2,704 municipalities '
      'have at least two seasons and therefore contribute at least one consecutive '
      'pair.')
corpo('Collecting from the primary source rather than from derived datasets is what '
      'makes states comparable: any third-party compilation imposes its own filters '
      'and time windows, so differences observed between units could reflect the '
      'compilation rather than the data. Here all municipalities are measured with '
      'the same instrument over the same period.')

sec('3.2. The Pará dataset, for the predictive analysis')
corpo('To examine the consequence of repetition for predictive models, we used a '
      'dataset built for the state of Pará, which integrates IBGE yield with NDVI and '
      'EVI vegetation indices from the MODIS sensor (MOD13Q1 product), precipitation '
      'from CHIRPS (Funk et al., 2015), and temperature, solar radiation and '
      'evapotranspiration from the ERA5-Land reanalysis (Muñoz-Sabater et al., 2021), '
      'aggregated over the window from November of the preceding year to May of the '
      'harvest year. Given the territorial extent of Amazonian municipalities, '
      'spectral indices were extracted exclusively from pixels classified as soybean '
      'in the annual MapBiomas mask (Collection 10.1, 30 m resolution; Souza et al., '
      '2020). This dataset comprises 415 records from 38 municipalities, with a mean '
      'yield of 2,995 kg ha⁻¹.')

sec('3.3. Definition and testing of the repetition rate')
corpo('A repetition is defined as a pair of consecutive seasons, within the same '
      'municipality, whose average yield values are exactly equal. The repetition '
      'rate is the ratio of repeated pairs to total observed pairs. The denominator '
      'is always the pair, never the record: a municipality with n seasons offers '
      'n−1 pairs, and dividing by the n records would dilute the rate unevenly across '
      'units, according to how many short series each contains.')
corpo('To assess whether the observed rate is compatible with chance, a permutation '
      'test was applied. Under the null hypothesis, the yield values of each '
      'municipality are randomly shuffled across years, which preserves the marginal '
      'distribution of each series — including any duplicated values — and destroys '
      'only the temporal ordering. The statistic is recomputed over 2,000 '
      'permutations.')
corpo('The null hypothesis is deliberately conservative. Shuffling within each '
      'municipality preserves all duplicates already present in the series: a '
      'municipality whose history contains three occurrences of 3,000 kg ha⁻¹ will '
      'still contain them after permutation, and those repetitions will keep '
      'producing identical pairs by chance. The rate expected under randomness is '
      'therefore not close to zero, and any excess observed above it cannot be '
      'attributed to rounding of the values. The descriptive significance level is '
      'reported as p < 0.0005, the lower bound imposed by the number of permutations.')

sec('3.4. Predictive models and baseline')
corpo('Four algorithms were trained — Random Forest (Breiman, 2001), XGBoost (Chen '
      'and Guestrin, 2016), Support Vector Regression and Multilayer Perceptron — '
      'under leave-one-year-out temporal cross-validation: each season is predicted '
      'by a model trained without it, which simulates the prediction of a future year '
      'and prevents temporal information leakage. Hyperparameters were set by nested '
      'search within each training partition, so that the result could not be '
      'attributed to insufficient tuning effort.')
corpo('The central element of the experimental design is the baseline. It predicts the '
      'yield of a municipality in a given season by that municipality\'s historical '
      'mean in the training set, plus the estimated temporal trend, and uses no '
      'climatic or spectral information. The comparison answers the question that '
      'actually matters: do environmental variables add predictive capacity beyond '
      'what the municipality\'s own history already conveys?')

# ══════════════════════════ 4 ══════════════════════════
sec('4. Results')

sec('4.1. Repetition in the Brazilian PAM')
corpo('Of the 46,536 consecutive-season pairs observed in Brazilian municipalities, '
      '8,192 — that is, 17.6% — report strictly identical yield. The rate is not '
      'uniformly distributed: it ranges from 3.7% in Piauí to 45.7% in Pará among '
      'units with sufficient sample, as Fig. 1 shows.')
cap('Fig. 1. Rate of consecutive seasons with identical soybean yield, by state. '
    'Units with fewer than 30 pairs are omitted. Pará is highlighted in dark red and '
    'São Paulo in orange. Source: table 5457, SIDRA/IBGE.')
corpo('Table 1 details the units with highest and lowest rates, with the number of '
      'pairs supporting each measure. Amazonas records the highest rate, 63.2%, but '
      'over only 38 pairs in 11 municipalities; among units with relevant production, '
      'Pará ranks first, with 45.7% over 492 pairs.')
cap('Table 1. Repetition rate by state (extremes and reference states).')
tabela(['State', 'Rate', 'Pairs', 'Municipalities', 'Multiples of 100'],
       [['Amazonas', '63.2%', '38', '11', '55.1%'],
        ['Pará', '45.7%', '492', '53', '77.4%'],
        ['Rondônia', '35.3%', '498', '46', '57.0%'],
        ['São Paulo', '31.9%', '7,762', '522', '69.5%'],
        ['Minas Gerais', '27.7%', '3,987', '322', '73.9%'],
        ['Goiás', '21.8%', '4,195', '240', '67.4%'],
        ['Mato Grosso', '12.4%', '2,479', '136', '43.7%'],
        ['Rio Grande do Sul', '11.5%', '9,303', '453', '57.9%'],
        ['Paraná', '5.2%', '8,522', '392', '40.8%'],
        ['Piauí', '3.7%', '482', '39', '22.3%'],
        ['Brazil', '17.6%', '46,536', '2,704', '58.4%']],
       negrito_em=(10,))
par('', dep=9)

sec('4.2. Repetition is not explained by chance')
corpo('An identical value between consecutive seasons could, in principle, occur by '
      'coincidence, particularly in series of rounded values. The permutation test '
      'rules this out with a wide margin. Under randomness the expected rate is 7.5% '
      '± 0.1%, and the maximum observed over 2,000 permutations was 7.9%. The rate '
      'actually recorded, 17.6%, lies more than one hundred standard deviations above '
      'the mean of the null distribution (p < 0.0005), as shown in Fig. 2. No '
      'permutation came close to the observed value.')
cap('Fig. 2. Permutation test: observed rate versus the distribution under '
    'randomness, obtained from 2,000 permutations of the municipal series.')

sec('4.3. Crop scale, not region, predicts repetition')
corpo('An immediate reading of Fig. 1 would suggest a regional effect: the leading '
      'positions are held by Amazonian states. The fourth position, however, belongs '
      'to São Paulo, with 31.9% over 7,762 pairs — a larger sample than that of any '
      'state in the North region, and in a unit that is not an agricultural frontier '
      'under any criterion. The regional hypothesis does not account for it.')
corpo('Crop scale does. Fig. 3 relates, for each municipality with at least five '
      'pairs, the soybean planted area and the observed repetition rate. The '
      'association is monotonic across four orders of magnitude: the mean per area '
      'decile falls from about 33% in the smallest-area decile to about 7% in the '
      'largest. The Spearman correlation between the logarithm of area and the rate '
      'is −0.459 (n = 2,460 municipalities; p < 10⁻¹²⁷).')
cap('Fig. 3. Municipal soybean planted area and repetition rate. Each point is a '
    'municipality with at least five consecutive-season pairs, coloured by whether '
    'it lies in one of the nine Legal Amazon states. Logarithmic scale on the '
    'abscissa.')
corpo('The decisive test is to exclude the Amazon. Restricting the analysis to the '
      '2,082 municipalities outside the Legal Amazon states, the correlation does not '
      'weaken: it strengthens to −0.503. Within the Legal Amazon itself, over 378 '
      'municipalities, it is −0.266. The gradient is therefore not an effect of the '
      'region; it is an effect of scale, which manifests equally — or more sharply — '
      'outside it.')
corpo('The pattern also holds within individual states. Dividing the municipalities of '
      'each state into quartiles of planted area, ten of the fourteen states with at '
      'least two hundred pairs show a strictly decreasing gradient, and in the '
      'remaining four the inversion occurs between the first and second quartiles, '
      'with the fourth always below the first. Table 2 presents selected cases.')
cap('Table 2. Repetition rate by quartile of planted area, within each state. Median '
    'area is per municipality-season; quartiles are computed within each state, with '
    'the consecutive-season pair as denominator.')
tabela(['State', 'Q1 (smallest)', 'Q2', 'Q3', 'Q4 (largest)', 'Median area'],
       [['São Paulo', '37.9%', '34.7%', '32.6%', '20.9%', '500 ha'],
        ['Minas Gerais', '34.8%', '30.4%', '24.3%', '21.3%', '1,200 ha'],
        ['Goiás', '34.5%', '23.2%', '17.5%', '11.7%', '—'],
        ['Rio Grande do Sul', '21.1%', '10.9%', '8.2%', '5.6%', '5,500 ha'],
        ['Mato Grosso', '18.1%', '15.6%', '9.2%', '6.8%', '27,912 ha'],
        ['Paraná', '7.5%', '5.2%', '4.9%', '3.3%', '8,562 ha'],
        ['Brazil', '30.5%', '20.1%', '12.0%', '7.4%', '—']],
       negrito_em=(6,))
par('', dep=9)
corpo('The interpretation is direct and consistent with the survey\'s declared '
      'methodology. The field agent estimates the municipal value from contacts with '
      'producers and technicians (IBGE, 2024). Where soybean occupies tens of '
      'thousands of hectares there are cooperatives, warehouses and technical '
      'assistance, and informants are abundant. Where the crop occupies a few hundred '
      'hectares — which occurs both in a recently opened Amazonian municipality and in '
      'a São Paulo municipality where soybean is a secondary crop — the informant base '
      'is sparse for that crop, and it is plausible that the previous season\'s value '
      'is carried forward in the absence of new information.')

sec('4.4. Structural patterns')
corpo('Two further characteristics reinforce this reading. First, repetition is not '
      'distributed as isolated noise but forms plateaus: 5,015 runs of consecutive '
      'years with identical values were identified, of which 363 extend for five years '
      'or more and 28 for ten years or more, reaching seventeen years in the extreme '
      'case.')
corpo('Second, value granularity is coarse: 58.4% of records are exact multiples of '
      '100 kg ha⁻¹ and 14.3% are multiples of 1,000. The value 3,000 kg ha⁻¹ alone '
      'accounts for 5,871 of the 49,342 records, that is, 11.9% of the entire national '
      'dataset. The five most frequent values — 3,000, 2,400, 2,700, 3,600 and 3,300 '
      'kg ha⁻¹ — are all multiples of 300, which corresponds to round numbers of '
      '60 kg bags per hectare. Fig. 4 illustrates the phenomenon in three municipal '
      'series.')
cap('Fig. 4. Municipal soybean yield series from the PAM/IBGE, for three '
    'municipalities in Pará, showing plateaus of constant value.')

sec('4.5. Consequence for predictive models')
corpo('Table 3 reports the performance of four algorithms in the state of Pará, '
      'against the baseline. None outperforms the baseline: the Multilayer Perceptron '
      'matches it (R² = 0.216), Support Vector Regression very nearly reaches it '
      '(0.208), and the tree-based methods degrade performance, behaviour typical of '
      'overfitting on small samples.')
cap('Table 3. Predictive performance in Pará under leave-one-year-out validation '
    '(2001–2024). The baseline uses no climatic or spectral variables. Dataset of 415 '
    'records, 38 municipalities, mean yield 2,995 kg ha⁻¹.')
tabela(['Model', 'RMSE (kg ha⁻¹)', 'MAE (kg ha⁻¹)', 'R²'],
       [['Baseline (history + trend)', '416', '296', '0.216'],
        ['Multilayer Perceptron', '416', '297', '0.216'],
        ['Support Vector Regression', '418', '295', '0.208'],
        ['Random Forest', '432', '309', '0.153'],
        ['XGBoost', '440', '315', '0.123']])
par('', dep=9)
corpo('The variance decomposition clarifies the result. Sixty-three per cent of the '
      'variability in Pará yields occurs within municipalities over time — precisely '
      'the component that environmental variables should explain. However, taking the '
      'deviation of each observation from its municipal mean, the correlations between '
      'that deviation and each predictor are negligible, not exceeding 0.20 in '
      'absolute value, which amounts to less than 4% of the within-municipality '
      'variation explained by any single predictor. If a substantial share of that '
      'interannual variation is not agronomic signal but administrative carry-forward '
      'of the previous value, no model can recover it.')

# ══════════════════════════ 5 ══════════════════════════
sec('5. Discussion')
corpo('These findings do not constitute a failure of the PAM, but a manifestation of '
      'its declared methodology operating under unequal conditions. IBGE states that '
      'municipal values are estimates produced by the field agent from contacts with '
      'producers and technicians and from local knowledge (IBGE, 2024). What this '
      'study adds is a measure of how much that condition varies, and the '
      'identification of what governs it.')
corpo('The regional hypothesis, which would be the natural reading, does not hold. If '
      'the problem belonged to the Amazon, the correlation between area and repetition '
      'would vanish when the Amazon is excluded; instead it strengthens, from −0.459 '
      'to −0.503. And São Paulo, with 31.9% over almost eight thousand pairs, would be '
      'inexplicable. The predictor is the scale of the crop in that municipality, not '
      'latitude: a five-hundred-hectare soybean crop has a sparse informant network in '
      'Ribeirão Preto as much as in Novo Progresso.')
corpo('The methodological implication is direct. A predictive model can only recover '
      'variation present in the target variable. If a fraction of interannual '
      'transitions is administratively constant, there is a ceiling on the attainable '
      'coefficient of determination, regardless of algorithm sophistication or '
      'predictor richness. This explains why, in Pará, performance saturates at the '
      'baseline level, whereas in Paraná — where repetition is 5.2% — the same '
      'algorithms and the same families of variables reach substantially higher R².')
corpo('It is worth being precise about what the ceiling means quantitatively. If a '
      'fraction of transitions is administratively constant, the observed interannual '
      'variance is smaller than the true agronomic variance, and the denominator of '
      'the coefficient of determination shrinks. The R² of 0.216 obtained in Pará '
      'should not, therefore, be read as weak performance of models that would achieve '
      'higher values elsewhere: it measures the explained fraction of a variance that '
      'has already been partly suppressed at source. Direct comparisons of R² between '
      'regions with differing survey quality are, in this sense, misleading, and the '
      'repetition rate offers a simple indicator with which to qualify them.')
corpo('This diagnosis speaks to two recent findings. Li et al. (2024) show that '
      'removing the global trend markedly improves the accuracy of XGBoost-based '
      'models, indicating that a relevant share of the variance explained in crop '
      'yield studies comes from the trend component of the series. Wang et al. (2024), '
      'confronting deep learning with panel regression in Argentina, find a smaller '
      'advantage than usually reported when the comparison is made against an explicit '
      'statistical reference. The difference is that in both studies the narrow margin '
      'follows from the statistical structure of the series; here it follows also from '
      'the way the series is produced.')
corpo('Three recommendations follow. First, crop yield prediction studies should report '
      'the repetition rate of the target variable, as one reports the proportion of '
      'missing data; the procedure is simple and reveals a limit that, ignored, is '
      'mistaken for model deficiency. Second, the selection of municipalities for '
      'predictive studies should treat crop scale as an explicit criterion, not merely '
      'data availability. Third, comparison against a baseline stripped of '
      'environmental variables should be standard practice: without it, a moderate R² '
      'can be read as predictive success when it reflects only the structure of '
      'municipal means.')
corpo('Limitations are acknowledged. The predictive analysis is restricted to Pará, so '
      'the consequence of repetition for models was measured in one state, even though '
      'the repetition diagnosis is national. Attributing the pattern to carry-forward '
      'by the field agent, although consistent with the official methodology and with '
      'the scale gradient, remains inferential: confirming it would require access to '
      'the survey\'s field records. And planted area, used as the measure of scale, '
      'comes from the same survey whose quality is under examination, although there '
      'is no reason to expect any error in it to be correlated with repetition in '
      'yield.')

# ══════════════════════════ 6 ══════════════════════════
sec('6. Conclusion')
corpo('This study documented that 17.6% of the 46,536 consecutive-season pairs of '
      'soybean in the Brazilian PAM report strictly identical yield — a rate lying '
      'more than one hundred standard deviations from that expected under randomness, '
      'with none of 2,000 permutations reaching it. The pattern forms plateaus of up '
      'to seventeen years and is associated with strong rounding, with 11.9% of the '
      'entire national dataset concentrated in a single value, 3,000 kg ha⁻¹.')
corpo('The central finding, however, is not the magnitude but the predictor. The '
      'repetition rate is not a function of region but of the scale of the crop in the '
      'municipality: the correlation between planted area and repetition is −0.459 '
      'nationally and −0.503 when the Legal Amazon states are excluded, and the '
      'gradient holds within ten of the fourteen states with sufficient sample. São '
      'Paulo, at 31.9%, demonstrates that the phenomenon is not Amazonian.')
corpo('As a consequence, machine learning models fed with climatic and spectral '
      'variables did not outperform, in Pará, a baseline built solely from municipal '
      'history and temporal trend. The limit lies not in the method but in the target '
      'variable. Beyond the diagnosis, this work provides a reproducible national '
      'dataset and a verification procedure applicable to other crops and territories, '
      'and recommends that the repetition rate be reported as an indicator of target '
      'variable quality in crop yield prediction studies.')

# ══════════════════════ DECLARATIONS ══════════════════════
sec('CRediT authorship contribution statement')
par('Maycon Lima dos Santos: Conceptualization, Methodology, Software, Formal '
    'analysis, Investigation, Data curation, Visualization, Validation, '
    'Writing – original draft, Writing – review and editing.',
    12, align=AL.JUSTIFY, lh=2.0, dep=12)

sec('Declaration of competing interest')
par('The author declares no known competing financial interests or personal '
    'relationships that could have appeared to influence the work reported in '
    'this paper.', 12, align=AL.JUSTIFY, lh=2.0, dep=12)

sec('Funding')
par('This research did not receive any specific grant from funding agencies in the '
    'public, commercial, or not-for-profit sectors.', 12, align=AL.JUSTIFY, lh=2.0, dep=12)

sec('Data availability')
par('The data and code that support the findings of this study are openly available. '
    'The national dataset assembled from SIDRA/IBGE table 5457 '
    '(pam_soja_municipios.csv, 49,342 records), the derived measures '
    '(repeticao_27_estados.json), the Pará dataset and all analysis scripts are '
    'archived at Zenodo under version v2.3.7, cited as dos Santos (2026), '
    'and are also available at '
    'https://github.com/engsoft7/dissertacao-soja-ia.', 12, align=AL.JUSTIFY, lh=2.0, dep=12)

sec('Declaration of generative AI and AI-assisted technologies in the manuscript '
    'preparation process')
par('During the preparation of this work the author used Claude (Anthropic) in order '
    'to assist with data analysis scripting, figure generation, English drafting and '
    'internal consistency checking of reported numerical values. After using this '
    'tool, the author reviewed and edited the content as needed and takes full '
    'responsibility for the content of the published article.',
    12, align=AL.JUSTIFY, lh=2.0, dep=12)

sec('Acknowledgements')
par('The author thanks the Graduate Program in Applied Computing of the Universidade '
    'Federal do Pará for institutional support.', 12, align=AL.JUSTIFY, lh=2.0, dep=12)

# ══════════════════════════ REFERENCES ══════════════════════════
sec('References')
refs = [
    'Barbosa dos Santos, V., Santos, A.M.F.D., Rolim, G.D.S., 2021. Estimation and '
    'forecasting of soybean yield using artificial neural networks. Agron. J. 113, '
    '3193–3209. https://doi.org/10.1002/agj2.20729.',
    'Breiman, L., 2001. Random forests. Mach. Learn. 45, 5–32.',
    'Chen, T., Guestrin, C., 2016. XGBoost: a scalable tree boosting system, in: '
    'Proceedings of the 22nd ACM SIGKDD International Conference on Knowledge Discovery '
    'and Data Mining. ACM, New York, pp. 785–794. '
    'https://doi.org/10.1145/2939672.2939785.',
    'dos Santos, M.L., 2026. Repetition in the Brazilian Municipal '
    'Agricultural Production survey: datasets and analysis code, v2.3.7 [dataset]. '
    'Zenodo. https://doi.org/10.5281/zenodo.22343315.',
    'Fathi, M., Shah-Hosseini, R., Moghimi, A., Arefi, H., 2025. '
    'MHRA-MS-3D-ResNet-BiLSTM: a multi-head-residual attention-based multi-stream deep '
    'learning model for soybean yield prediction in the U.S. using multi-source remote '
    'sensing data. Remote Sens. 17, 107. https://doi.org/10.3390/rs17010107.',
    'Fu, H., Li, J., Lu, J., Lin, X., Kang, J., Zou, W., Ning, X., Sun, Y., 2025. '
    'Prediction of soybean yield at the county scale based on multi-source '
    'remote-sensing data and deep learning models. Agriculture 15, 1337. '
    'https://doi.org/10.3390/agriculture15131337.',
    'Funk, C., Peterson, P., Landsfeld, M., Pedreros, D., Verdin, J., Shukla, S., '
    'Husak, G., Rowland, J., Harrison, L., Hoell, A., Michaelsen, J., 2015. The climate '
    'hazards infrared precipitation with stations: a new environmental record for '
    'monitoring extremes. Sci. Data 2, 150066. https://doi.org/10.1038/sdata.2015.66.',
    'IBGE, 2024. Produção Agrícola Municipal: notas metodológicas. Instituto Brasileiro '
    'de Geografia e Estatística, Rio de Janeiro.',
    'Ingole, V.S., Kshirsagar, U.A., Singh, V., Yadav, M.V., Krishna, B., Kumar, R., '
    '2025. A hybrid model for soybean yield prediction integrating convolutional neural '
    'networks, recurrent neural networks, and graph convolutional networks. Computation '
    '13, 4. https://doi.org/10.3390/computation13010004.',
    'Khaki, S., Wang, L., 2019. Crop yield prediction using deep neural networks. '
    'Front. Plant Sci. 10, 621. https://doi.org/10.3389/fpls.2019.00621.',
    'Khaki, S., Wang, L., Archontoulis, S.V., 2020. A CNN-RNN framework for crop yield '
    'prediction. Front. Plant Sci. 10, 1750. https://doi.org/10.3389/fpls.2019.01750.',
    'Leukel, J., Zimpel, T., Stumpe, C., 2023. Machine learning technology for early '
    'prediction of grain yield at the field scale: a systematic review. Comput. '
    'Electron. Agric. 207, 107721. https://doi.org/10.1016/j.compag.2023.107721.',
    'Li, Y., Zeng, H., Zhang, M., Wu, B., Qin, X., 2024. Global de-trending '
    'significantly improves the accuracy of XGBoost-based county-level maize and '
    'soybean yield prediction in the Midwestern United States. GIsci. Remote Sens. 61, '
    '2349341. https://doi.org/10.1080/15481603.2024.2349341.',
    'Maimaitijiang, M., Sagan, V., Sidike, P., Hartling, S., Esposito, F., Fritschi, '
    'F.B., 2020. Soybean yield prediction from UAV using multimodal data fusion and '
    'deep learning. Remote Sens. Environ. 237, 111599. '
    'https://doi.org/10.1016/j.rse.2019.111599.',
    'Muñoz-Sabater, J., Dutra, E., Agustí-Panareda, A., Albergel, C., Arduini, G., '
    'Balsamo, G., Boussetta, S., Choulga, M., Harrigan, S., Hersbach, H., Martens, B., '
    'Miralles, D.G., Piles, M., Rodríguez-Fernández, N.J., Zsoter, E., Buontempo, C., '
    'Thépaut, J.-N., 2021. ERA5-Land: a state-of-the-art global reanalysis dataset for '
    'land applications. Earth Syst. Sci. Data 13, 4349–4383.',
    'Richetti, J., Judge, J., Boote, K.J., Johann, J.A., Uribe-Opazo, M.A., Becker, '
    'W.R., Paludo, A., Silva, L.C.d.A., 2018. Using phenology-based enhanced vegetation '
    'index and machine learning for soybean yield estimation in Paraná State, Brazil. '
    'J. Appl. Remote Sens. 12, 026029. https://doi.org/10.1117/1.JRS.12.026029.',
    'Schwalbert, R.A., Amado, T., Corassa, G., Pott, L.P., Prasad, P.V.V., Ciampitti, '
    'I.A., 2020. Satellite-based soybean yield forecast: integrating machine learning '
    'and weather data for improving crop yield prediction in southern Brazil. Agric. '
    'For. Meteorol. 284, 107886. https://doi.org/10.1016/j.agrformet.2019.107886.',
    'Souza, C.M., Shimbo, J.Z., Rosa, M.R., Parente, L.L., Alencar, A.A., Rudorff, '
    'B.F.T., Hasenack, H., Matsumoto, M., Ferreira, L.G., Souza-Filho, P.W.M., de '
    'Oliveira, S.W., Rocha, W.F., Fonseca, A.V., Marques, C.B., Diniz, C.G., Costa, D., '
    'Monteiro, D., Rosa, E.R., Vélez-Martin, E., Weber, E.J., Lenti, F.E.B., Paternost, '
    'F.F., Pareyn, F.G.C., Siqueira, J.V., Viera, J.L., Ferreira Neto, L.C., Saraiva, '
    'M.M., Sales, M.H., Salgado, M.P.G., Vasconcelos, R., Galano, S., Mesquita, V.V., '
    'Azevedo, T., 2020. Reconstructing three decades of land use and land cover changes '
    'in Brazilian biomes with Landsat archive and Earth Engine. Remote Sens. 12, 2735. '
    'https://doi.org/10.3390/rs12172735.',
    'Sun, J., Di, L., Sun, Z., Shen, Y., Lai, Z., 2019. County-level soybean yield '
    'prediction using deep CNN-LSTM model. Sensors 19, 4363. '
    'https://doi.org/10.3390/s19204363.',
    'van Klompenburg, T., Kassahun, A., Catal, C., 2020. Crop yield prediction using '
    'machine learning: a systematic literature review. Comput. Electron. Agric. 177, '
    '105709.',
    'von Bloh, M., Nóia Júnior, R.d.S., Wangerpohl, X., Saltık, A.O., Haller, V., '
    'Kaiser, L., Asseng, S., 2023. Machine learning for soybean yield forecasting in '
    'Brazil. Agric. For. Meteorol. 341, 109670. '
    'https://doi.org/10.1016/j.agrformet.2023.109670.',
    'Wang, Y., Feng, K., Sun, L., Xie, Y., Song, X.P., 2024. Satellite-based soybean '
    'yield prediction in Argentina: a comparison between panel regression and deep '
    'learning methods. Comput. Electron. Agric. 221, 108978. '
    'https://doi.org/10.1016/j.compag.2024.108978.',
    'You, J., Li, X., Low, M., Lobell, D., Ermon, S., 2017. Deep Gaussian process for '
    'crop yield prediction based on remote sensing data, in: Proceedings of the 31st '
    'AAAI Conference on Artificial Intelligence. AAAI Press, Palo Alto, pp. 4559–4566. '
    'https://doi.org/10.1609/aaai.v31i1.11172.',
]
for r in refs:
    par(r, 12, align=AL.JUSTIFY, lh=1.5, dep=6)

d.save(os.path.join(SAIDA, 'Manuscript_CEA.docx'))

# ─────────── Highlights, arquivo separado ───────────
h = novo_doc()
par('Highlights', 12, True, align=AL.LEFT, lh=1.0, dep=12, doc=h)
for b in [
    '17.6% of consecutive soybean seasons in Brazil report identical yield',
    'Observed rate exceeds chance by more than 100 standard deviations',
    'Crop scale, not region, predicts repetition (Spearman −0.46; −0.50 outside Amazon)',
    'São Paulo reaches 31.9%, showing the effect is not an Amazonian phenomenon',
    'Environmental predictors do not beat a municipal history-plus-trend baseline',
]:
    assert len(b) <= 85, (len(b), b)
    par('• ' + b, 12, align=AL.LEFT, lh=1.5, dep=6, doc=h)
h.save(os.path.join(SAIDA, 'Highlights.docx'))
print('gerados')
